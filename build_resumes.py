# Builds Zain Iqbal's three tailored resumes as ATS-friendly DOCX files.
# ASCII separators only (| and -). Run: python build_resumes.py
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATE = "2026-05-29"
OUTDIR = r"d:\Claude\iZainIqbal"
ACCENT = RGBColor(0x1F, 0x49, 0x7D); DARK = RGBColor(0x22, 0x22, 0x22); GREY = RGBColor(0x55, 0x55, 0x55)

CONTACT = [("LinkedIn", "https://www.linkedin.com/in/zain-iqbal-devs/"),
           ("Portfolio", "https://izainiqbal.github.io/my-portfolio/"),
           ("GitHub", "https://github.com/iZainIqbal")]
PHONE_EMAIL = "Lahore, Pakistan  |  +92 346 6088504  |  appdev.zain@gmail.com"
HANDMAN_LINKS = [("handman.ch", "https://handman.ch"), ("ch.handyman.app", "https://play.google.com/store/apps/details?id=ch.handyman.app")]
CS_AOE = [("Case Study", "https://metavizai.com/case-study/wellness-accountability-app-a-complete-wellness-platform-by-metaviz-ai/")]
CS_AIAPEX = [("Case Study", "https://metavizai.com/case-study/ai-apex-personalized-fitness-nutrition-coaching-powered-by-generative-ai/")]
CS_RAWTEEN = [("Case Study", "https://metavizai.com/case-study/woocommerce-based-grocery-delivery-app-nutrition-focused-e-commerce-for-rawteen-dubai/")]


def add_hyperlink(p, text, url, size=9):
    r_id = p.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1F497D"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rpr.append(sz)
    run.append(rpr); t = OxmlElement("w:t"); t.text = text; run.append(t)
    h.append(run); p._p.append(h)


def section_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:space"), "1"); b.set(qn("w:color"), "1F497D")
    pBdr.append(b); pPr.append(pBdr)


def dated_line(doc, left, right, left_size=10.5, before=4, after=1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    rl = p.add_run(left); rl.bold = True; rl.font.size = Pt(left_size)
    if right:
        rr = p.add_run("\t" + right); rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = GREY


def note_line(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.space_after = Pt(1.5)
    p.add_run(text).font.size = Pt(10)


def cert_bullet(doc, text, url=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.space_after = Pt(1.5)
    p.add_run(text).font.size = Pt(10)
    if url:
        s = p.add_run("  |  "); s.font.size = Pt(9); s.font.color.rgb = GREY
        add_hyperlink(p, "Verify", url)


def links_line(doc, pairs):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Links: "); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    for i, (text, url) in enumerate(pairs):
        if i:
            s = p.add_run("  |  "); s.font.size = Pt(9); s.font.color.rgb = GREY
        add_hyperlink(p, text, url)


def plain(doc, runs, size=10, after=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    for text, b in runs:
        r = p.add_run(text); r.bold = b; r.font.size = Pt(size)


def build(filename, role_title, summary, competencies, experience, projects, certs):
    doc = Document(); s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(0.4); s.left_margin = s.right_margin = Inches(0.6)
    base = doc.styles["Normal"].font; base.name = "Calibri"; base.size = Pt(10)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ZAIN IQBAL"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ACCENT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(role_title); r.font.size = Pt(12); r.font.color.rgb = DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    p.add_run(PHONE_EMAIL).font.size = Pt(9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    for i, (label, url) in enumerate(CONTACT):
        if i: p.add_run("  |  ").font.size = Pt(9)
        add_hyperlink(p, label, url)

    section_heading(doc, "Summary"); plain(doc, [(summary, False)])
    section_heading(doc, "Core Competencies")
    for label, items in competencies:
        plain(doc, [(label + ": ", True), (items, False)], after=1)

    section_heading(doc, "Professional Experience")
    for role in experience:
        dated_line(doc, role["role"], role["date"])
        if role.get("note"): note_line(doc, role["note"])
        for proj in role.get("projects", []):
            dated_line(doc, proj["name"], proj["date"], left_size=10, before=3, after=1)
            for b in proj["bullets"]:
                bullet(doc, b)
            if proj.get("links"): links_line(doc, proj["links"])

    if projects:
        section_heading(doc, "Projects")
        for proj in projects:
            dated_line(doc, proj["name"], proj.get("date", ""), left_size=10, before=3, after=1)
            for b in proj["bullets"]:
                bullet(doc, b)
            if proj.get("links"): links_line(doc, proj["links"])

    section_heading(doc, "Education")
    dated_line(doc, "BS Computer Science | University of Gujrat", "Nov 2021 - May 2025")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run("GPA 3.37 / 4.00. Key courses: OOP, Data Structures, Algorithms, DBMS, ML, AI, Web Development, Software Engineering.").font.size = Pt(10)

    section_heading(doc, "Certifications / Leadership")
    for text, url in certs:
        cert_bullet(doc, text, url)

    path = os.path.join(OUTDIR, filename); doc.save(path); print("wrote", path)


# ---------------- certifications ----------------
FIN = ("Finance Secretary | Hayatian Computing Society, University of Gujrat", None)
META_FE = ("Introduction to Front-End Development | Meta (2023)", "https://www.coursera.org/account/accomplishments/verify/ZVN4DA64FZHZ")
META_HTML = ("HTML and CSS in Depth | Meta (2023)", "https://www.coursera.org/account/accomplishments/verify/BMSS7PNSRBB8")
UNITY = ("Fundamentals of Unity Android Game Development | Packt (2025)", "https://www.coursera.org/account/accomplishments/verify/78FKGBXM720Z")
PROBSOLVE = ("Effective Problem-Solving and Decision-Making | UC Irvine (2025)", "https://www.coursera.org/account/accomplishments/verify/YL0U6EBSL387")
flutter_certs = [META_FE, UNITY, PROBSOLVE, FIN]
fullstack_certs = [META_FE, META_HTML, PROBSOLVE, FIN]
combined_certs = [META_FE, META_HTML, UNITY, PROBSOLVE, FIN]

NOTE = "Joined as an intern (Aug 2025) and converted to a full-time developer."

snapshop_proj = {"name": "Snap & Shop | E-commerce Web App (Metaviz AI)", "bullets": [
    "Built a full-stack e-commerce web app: React frontend with a Node.js/Express + MongoDB backend.",
    "Designed REST APIs and MongoDB schemas for products, cart, orders, and users; integrated Stripe payments and Google OAuth.",
]}
chrono_proj = {"name": "Chrono Chase | Unity (Final Year Project)", "bullets": [
    "Mobile endless-runner with a time-travel theme, in-game shop, and modular gameplay logic.",
]}

# ---------------- 1. FLUTTER ----------------
build(
    f"Zain_Iqbal_Resume_Flutter_{DATE}.docx",
    "Flutter Developer",
    "Flutter Developer building production cross-platform Android and iOS apps with Firebase, REST APIs, and payment integrations. Built the complete frontend of a 100+ screen Swiss marketplace published on Google Play, and delivered two full-stack AI and wellness apps end-to-end. Strong in clean-architecture mobile design, payment systems, real-time features, and multilingual apps.",
    [
        ("Mobile", "Flutter, Dart, Android, iOS, Clean Architecture, BLoC, Riverpod, Provider, Dependency Injection (get_it), SQLite (sqflite), Offline Support"),
        ("Firebase", "Authentication, Firestore, Realtime Database, Cloud Messaging (FCM)"),
        ("APIs & Integrations", "REST APIs, FastAPI, SSE, Stripe, TWINT, PostFinance, Twilio OTP, Google Maps API, Google OAuth, Push Notifications"),
        ("Tools", "Git, Android Studio, VS Code, Google Play Console, Firebase Console, Docker, Coolify"),
        ("Languages", "Dart, JavaScript, TypeScript, Python, SQL"),
    ],
    [
        {"role": "Flutter Developer | Metaviz AI", "date": "Aug 2025 - Present  |  Remote", "note": NOTE, "projects": [
            {"name": "Handman (Switzerland) | Swiss Marketplace Platform", "date": "Mar 2026 - Present", "bullets": [
                "Built the complete Flutter frontend (100+ screens, 4 user roles) for a Swiss craftsmen marketplace, published on Google Play; clean-architecture codebase with offline support (sqflite). (team)",
                "Built real-time chat, a quotes-to-invoice-to-billing flow, and an AI assistant, with a 4-language UI (German, French, Italian, English) on Firebase Auth, Realtime DB, and FCM.",
                "Integrated Stripe, TWINT, and PostFinance payments, Twilio OTP, and biometric/passkey login, with SSE live updates from the FastAPI backend.",
            ], "links": HANDMAN_LINKS},
            {"name": "RawTeen (Dubai) | Grocery Delivery App", "date": "Jan - Feb 2026", "bullets": [
                "Built the Flutter frontend for a Dubai grocery-delivery app and connected it to the backend.",
                "Integrated Stripe checkout, a WooCommerce REST catalog, and a Google Maps store locator with delivery zones.",
            ], "links": CS_RAWTEEN},
            {"name": "AOE Wellness | Wellness Platform", "date": "Nov - Dec 2025", "bullets": [
                "Built end-to-end (frontend + backend): gamified habit-tracking and role-based Member/Distributor dashboards.",
            ], "links": CS_AOE},
            {"name": "AI-APEX | AI Fitness & Nutrition App", "date": "Sep - Oct 2025", "bullets": [
                "Sole developer (frontend + backend) of an AI fitness app; integrated the Claude API, Firebase, and Stripe subscriptions.",
            ], "links": CS_AIAPEX},
        ]},
    ],
    [],
    flutter_certs,
)

# ---------------- 2. FULL STACK ----------------
build(
    f"Zain_Iqbal_Resume_FullStack_{DATE}.docx",
    "Full Stack Developer",
    "Full Stack Developer building production web and mobile apps with React, TypeScript, FastAPI, Node.js, Express, MongoDB, and PostgreSQL. Built two full-stack AI and wellness apps end-to-end, plus the web frontend and payment/real-time backend of a live 100+ screen Swiss marketplace, deployed with Docker on Coolify. Strong in REST API design, payment systems, and containerized deployment.",
    [
        ("Frontend", "React, TypeScript, JavaScript, Tailwind CSS, Vite, TanStack Query, React Hook Form"),
        ("Backend", "FastAPI (Python), Node.js, Express, REST APIs, Server-Sent Events (SSE)"),
        ("Databases", "MongoDB, PostgreSQL, Firebase (Firestore, Realtime DB), SQLite"),
        ("Deployment & Infra", "Docker, Coolify, Redis, Celery, background workers"),
        ("Integrations", "Stripe, Stripe Connect, TWINT, PostFinance, Twilio OTP, Webhooks, Google OAuth, Firebase Auth, FCM"),
        ("Mobile", "Flutter (Android, iOS), Dart"),
        ("Familiar / learning", "Next.js, Nest.js, T3 stack"),
        ("Languages", "JavaScript, TypeScript, Python, Dart, SQL"),
    ],
    [
        {"role": "Full Stack Developer | Metaviz AI", "date": "Aug 2025 - Present  |  Remote", "note": NOTE, "projects": [
            {"name": "Handman (Switzerland) | Swiss Marketplace Platform", "date": "Mar 2026 - Present", "bullets": [
                "Built the complete React 19 + TypeScript (Vite) web frontend for a 100+ screen, 4-role Swiss marketplace (Tailwind, TanStack Query, 4 languages), live at handman.ch.",
                "Developed FastAPI backend endpoints for payments (Stripe, TWINT, PostFinance), Twilio OTP verification, and SSE real-time updates.",
                "Created Docker images and deployed services on Coolify; ran Redis-backed Celery workers for background jobs.",
            ], "links": HANDMAN_LINKS},
            {"name": "RawTeen (Dubai) | Grocery Delivery App", "date": "Jan - Feb 2026", "bullets": [
                "Built the frontend and connected the backend; integrated Stripe and a WooCommerce REST catalog with a Google Maps store locator.",
            ], "links": CS_RAWTEEN},
            {"name": "AOE Wellness | Wellness Platform", "date": "Nov - Dec 2025", "bullets": [
                "Built end-to-end (frontend + backend): role-based Member/Distributor dashboards, REST APIs, and content delivery.",
            ], "links": CS_AOE},
            {"name": "AI-APEX | AI Fitness & Nutrition App", "date": "Sep - Oct 2025", "bullets": [
                "Sole developer (frontend + backend): Node.js/Python APIs + Flutter client; integrated the Claude API, Firebase, and Stripe subscriptions.",
            ], "links": CS_AIAPEX},
        ]},
    ],
    [snapshop_proj],
    fullstack_certs,
)

# ---------------- 3. COMBINED ----------------
build(
    f"Zain_Iqbal_Resume_Combined_{DATE}.docx",
    "Full Stack Engineer (Flutter & Web)",
    "Full Stack Engineer shipping production mobile and web apps for international clients. Built two full-stack AI and wellness apps end-to-end, plus the complete frontend (web + mobile) and payment/real-time backend of a live 100+ screen Swiss marketplace, deployed with Docker on Coolify. Stack: Flutter, React/TypeScript, FastAPI, Node.js/MongoDB. Strong in payment systems, real-time features, and containerized deployment.",
    [
        ("Mobile", "Flutter, Dart, Android, iOS, Clean Architecture, BLoC, Riverpod, Provider, get_it (DI), SQLite (sqflite)"),
        ("Frontend", "React, TypeScript, Tailwind CSS, Vite, TanStack Query"),
        ("Backend", "FastAPI (Python), Node.js, Express, REST APIs, SSE"),
        ("Databases & Infra", "MongoDB, PostgreSQL, Firebase, SQLite, Docker, Coolify, Redis, Celery"),
        ("Integrations", "Stripe, Stripe Connect, TWINT, PostFinance, Twilio OTP, Webhooks, Google Maps API, Google OAuth, Firebase Auth, FCM, Claude API"),
        ("Familiar / learning", "Next.js, Nest.js, T3 stack"),
        ("Languages", "Dart, JavaScript, TypeScript, Python, SQL"),
    ],
    [
        {"role": "Full Stack Engineer | Metaviz AI", "date": "Aug 2025 - Present  |  Remote", "note": NOTE, "projects": [
            {"name": "Handman (Switzerland) | Swiss Marketplace Platform", "date": "Mar 2026 - Present", "bullets": [
                "Built the complete frontend - web (React 19 + TypeScript, Tailwind) and mobile (Flutter, 100+ screens, 4 roles, clean architecture) - for a Swiss marketplace, live at handman.ch and on Google Play. (team)",
                "Developed FastAPI backend endpoints for payments (Stripe, TWINT, PostFinance), Twilio OTP, and SSE real-time updates.",
                "Created Docker images and deployed services on Coolify with Redis-backed Celery workers for background jobs.",
            ], "links": HANDMAN_LINKS},
            {"name": "RawTeen (Dubai) | Grocery Delivery App", "date": "Jan - Feb 2026", "bullets": [
                "Built the Flutter frontend, connected the backend, and integrated Stripe checkout and a WooCommerce REST catalog with a Google Maps store locator.",
            ], "links": CS_RAWTEEN},
            {"name": "AOE Wellness | Wellness Platform", "date": "Nov - Dec 2025", "bullets": [
                "Built end-to-end (frontend + backend): gamified habit tracking, role-based Member/Distributor dashboards, and REST APIs.",
            ], "links": CS_AOE},
            {"name": "AI-APEX | AI Fitness & Nutrition App", "date": "Sep - Oct 2025", "bullets": [
                "Sole developer (frontend + backend): Flutter client + Node.js/Python APIs; integrated the Claude API, Firebase, and Stripe subscriptions.",
            ], "links": CS_AIAPEX},
        ]},
    ],
    [snapshop_proj, chrono_proj],
    combined_certs,
)

print("Done.")
